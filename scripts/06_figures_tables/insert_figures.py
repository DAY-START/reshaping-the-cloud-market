# -*- coding: utf-8 -*-

# ============================================================================
# RESTRICTED USE -- All rights reserved.
# The data, experimental results, and source code in this repository are
# provided for archival / reproducibility reference only. No part of the
# data, experiments, or code may be used, copied, modified, redistributed,
# or incorporated into other works without the author's explicit written
# permission. See LICENSE and NOTICE.md.
# ============================================================================

"""Insert the regenerated publication-quality figures into V7 and save V8."""
import os 
from docx import Document 
from docx .shared import Inches ,Pt 
from docx .text .paragraph import Paragraph 
from docx .oxml import OxmlElement 

DOC_IN =r"C:\Users\ADMIN\Desktop\重塑云市场_AI算力关注与股价分化_V7_顶刊批注整合版.docx"
DOC_OUT =r"C:\Users\ADMIN\Desktop\重塑云市场_AI算力关注与股价分化_V8_顶刊批注整合版_含图.docx"
FIG_DIR =r"C:\Users\ADMIN\WorkBuddy\2026-08-02-01-36-28\fig_v8"

def insert_picture_before (paragraph ,img_path ,width_inches =6.0 ):
    p =paragraph ._p 
    new_p =OxmlElement ('w:p')
    p .addprevious (new_p )
    new_para =Paragraph (new_p ,paragraph ._parent )
    run =new_para .add_run ()
    run .add_picture (img_path ,width =Inches (width_inches ))
    new_para .alignment =1 # center
    return new_para 

def insert_caption_before (paragraph ,text ,italic =True ,size =10 ):
    p =paragraph ._p 
    new_p =OxmlElement ('w:p')
    p .addprevious (new_p )
    new_para =Paragraph (new_p ,paragraph ._parent )
    new_para .alignment =1 
    run =new_para .add_run (text )
    run .font .name ="Arial"
    run .font .size =Pt (size )
    if italic :
        run .italic =True 
    return new_para 

def insert_picture_after (paragraph ,img_path ,width_inches =6.0 ):
    p =paragraph ._p 
    new_p =OxmlElement ('w:p')
    p .addnext (new_p )
    new_para =Paragraph (new_p ,paragraph ._parent )
    run =new_para .add_run ()
    run .add_picture (img_path ,width =Inches (width_inches ))
    new_para .alignment =1 
    return new_para 

def insert_caption_after (paragraph ,text ,italic =True ,size =10 ):
    p =paragraph ._p 
    new_p =OxmlElement ('w:p')
    p .addnext (new_p )
    new_para =Paragraph (new_p ,paragraph ._parent )
    new_para .alignment =1 # center
    run =new_para .add_run (text )
    run .font .name ="Arial"
    run .font .size =Pt (size )
    if italic :
        run .italic =True 
    return new_para 

def find_para (doc ,predicate ):
    for p in doc .paragraphs :
        if predicate (p .text ):
            return p 
    return None 

def find_table (doc ,first_cell_predicate ):
    for t in doc .tables :
        if t .rows and first_cell_predicate (t .rows [0 ].cells [0 ].text ):
            return t 
    return None 

    # -----------------------------------------------------------------
doc =Document (DOC_IN )

# Remove any pre-existing low-quality figures so that only the regenerated ones remain
W ="{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
removed =0 
for p in list (doc .paragraphs ):
    if p ._p .find (".//"+W +"drawing")is not None :
        p ._p .getparent ().remove (p ._p )
        removed +=1 
print ("Removed pre-existing images:",removed )

# 1) Figure 1 — existing caption placeholder in §1
p =find_para (doc ,lambda t :"Figure 1. Research Framework"in t )
if p :
    insert_picture_before (p ,os .path .join (FIG_DIR ,"Figure1_framework.png"),width_inches =6.0 )
    print ("Inserted Figure1")
else :
    print ("WARN: Figure1 caption not found")

    # 2) Figure 2 — existing caption placeholder in §3
p =find_para (doc ,lambda t :"Figure 2. Data Processing"in t )
if p :
    insert_picture_before (p ,os .path .join (FIG_DIR ,"Figure2_workflow.png"),width_inches =6.2 )
    print ("Inserted Figure2")
else :
    print ("WARN: Figure2 caption not found")

    # 3) Figures 3 & 4 before Table 8 (first result table)
p8 =find_para (doc ,lambda t :t .strip ().startswith ("Table 8. Ablation experiments"))
if p8 :
# Build bottom-up before p8: Fig4 caption, Fig4 image, Fig3 caption, Fig3 image
    cap4 =insert_caption_before (p8 ,"Figure 4. Monthly AI-Compute News Composition by Theme")
    p4 =insert_picture_before (cap4 ,os .path .join (FIG_DIR ,"Figure4_theme_composition.png"),width_inches =6.3 )
    cap3 =insert_caption_before (p4 ,"Figure 3. Daily General Computing-concern Sentiment (GCS) with 30-Day Moving Average and Major Events")
    p3 =insert_picture_before (cap3 ,os .path .join (FIG_DIR ,"Figure3_daily_gcs_ma_events.png"),width_inches =6.5 )
    print ("Inserted Figure3/4 before Table8")
else :
    print ("WARN: Table8 not found")

    # 4) Figure 6 before Table 10 (nonlinear)
p10 =find_para (doc ,lambda t :t .strip ().startswith ("Table 10. Nonlinear effects"))
if p10 :
    cap6 =insert_caption_before (p10 ,"Figure 6. Conditional GPU-CPU Return Spread across UGCS Percentiles")
    insert_picture_before (cap6 ,os .path .join (FIG_DIR ,"Figure6_quantile_spread.png"),width_inches =6.0 )
    print ("Inserted Figure6 before Table10")
else :
    print ("WARN: Table10 not found")

    # 5) Figure 7 after the LP paragraph (contains formula (24) is the local projection)
lp_para =find_para (doc ,lambda t :"Equation (24) is the local projection"in t )
if lp_para :
    p7 =insert_picture_after (lp_para ,os .path .join (FIG_DIR ,"Figure7_local_projection.png"),width_inches =6.4 )
    insert_caption_after (p7 ,"Figure 7. Daily, Weekly and Monthly GPU-CPU Return Responses")
    print ("Inserted Figure7 after LP paragraph")
else :
    print ("WARN: LP paragraph not found")

    # 6) Figure 5 before Table 11 (event study)
p11 =find_para (doc ,lambda t :t .strip ().startswith ("Table 11. Event Study"))
if p11 :
    cap5 =insert_caption_before (p11 ,"Figure 5. Cumulative Abnormal Returns around Extreme UGCS Events (top-10%)")
    insert_picture_before (cap5 ,os .path .join (FIG_DIR ,"Figure5_event_study_car.png"),width_inches =6.0 )
    print ("Inserted Figure5 before Table11")
else :
    print ("WARN: Table11 not found")

    # 7) Figure A1 & A2 before Appendix B heading (inside Appendix A)
pB =find_para (doc ,lambda t :t .strip ().startswith ("Appendix B."))
if pB :
    capA2 =insert_caption_before (pB ,"Figure A2. DA-MT-FinTransformer Confusion Matrices (Full model)")
    pA2 =insert_picture_before (capA2 ,os .path .join (FIG_DIR ,"FigureA2_confusion.png"),width_inches =6.3 )
    capA1 =insert_caption_before (pA2 ,"Figure A1. Topic Correlation Heatmap and Hierarchical Clustering")
    pA1 =insert_picture_before (capA1 ,os .path .join (FIG_DIR ,"FigureA1_topic_corr.png"),width_inches =5.8 )
    print ("Inserted FigureA1/A2 before Appendix B")
else :
    print ("WARN: Appendix B heading not found")

    # save
doc .save (DOC_OUT )
print ("Saved",DOC_OUT )
