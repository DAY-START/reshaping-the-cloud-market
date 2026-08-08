# -*- coding: utf-8 -*-

# ============================================================================
# RESTRICTED USE -- All rights reserved.
# The data, experimental results, and source code in this repository are
# provided for archival / reproducibility reference only. No part of the
# data, experiments, or code may be used, copied, modified, redistributed,
# or incorporated into other works without the author's explicit written
# permission. See LICENSE and NOTICE.md.
# ============================================================================

"""Scan the V6 paper docx structure: title level, placeholders, table positions, and prepare for result filling."""
import re 
import sys 
import io 
from docx import Document 

sys .stdout =io .TextIOWrapper (sys .stdout .buffer ,encoding ="utf-8",errors ="replace")

SRC =r"C:\Users\ADMIN\Desktop\重塑云市场_AI算力替代如何驱动股价分化_V6.docx"

d =Document (SRC )
print (f"paragraphs={len (d .paragraphs )}  tables={len (d .tables )}")

print ("\n===== 标题结构 =====")
for i ,p in enumerate (d .paragraphs ):
    st =(p .style .name or "")
    t =p .text .strip ()
    if not t :
        continue 
    if st .startswith ("Heading")or re .match (r"^(\d+[\.、]|[一二三四五六七八九十]+[、\.]|附录|参考文献|摘要|Abstract)",t ):
        print (f"[{i :4d}] ({st }) {t [:80 ]}")

print ("\n===== 占位符 =====")
pat =re .compile (r"【[^】]{0,80}】|\[待补[^\]]*\]|XXX|待填|TBD")
cnt =0 
for i ,p in enumerate (d .paragraphs ):
    for m in pat .findall (p .text ):
        print (f"[{i :4d}] {m }")
        cnt +=1 
print (f"Total{cnt } places")

print ("\n===== 表格概览 =====")
for ti ,tb in enumerate (d .tables ):
    hdr =" | ".join (c .text .strip ()[:18 ]for c in tb .rows [0 ].cells )
    print (f"T{ti }: {len (tb .rows )}row x{len (tb .columns )}list header: {hdr [:120 ]}")
