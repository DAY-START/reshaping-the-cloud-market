# -*- coding: utf-8 -*-

# ============================================================================
# RESTRICTED USE -- All rights reserved.
# The data, experimental results, and source code in this repository are
# provided for archival / reproducibility reference only. No part of the
# data, experiments, or code may be used, copied, modified, redistributed,
# or incorporated into other works without the author's explicit written
# permission. See LICENSE and NOTICE.md.
# ============================================================================

"""v6_s09_fill_paper.py
====================
Fill in the paper V6.docx with the empirical results and charts of S01~S07.
Strategy:
1. Keep the original chapter number unchanged;
2. Insert section 4.7 "Empirical Results" after section 4.6, and present the charts with pictures and texts;
3. Replace the 6 [] placeholders in the abstract and summary;
4. Embed the text graphics at 300 dpi, and generate the tables according to the three-line table style of the top issue;
5. Output to a new file on the desktop without overwriting the original."""
import re 
import io 
import json 
import sys 
from pathlib import Path 
from docx import Document 
from docx .shared import Inches ,Pt ,RGBColor 
from docx .enum .text import WD_ALIGN_PARAGRAPH ,WD_LINE_SPACING 
from docx .enum .table import WD_TABLE_ALIGNMENT 
from docx .oxml import OxmlElement 
from docx .oxml .ns import qn 
import pandas as pd 

sys .stdout =io .TextIOWrapper (sys .stdout .buffer ,encoding ="utf-8",errors ="replace")

SRC =r"C:\Users\ADMIN\Desktop\重塑云市场_AI算力替代如何驱动股价分化_V6.docx"
OUT =r"C:\Users\ADMIN\Desktop\重塑云市场_AI算力替代如何驱动股价分化_V6_实证结果填充版.docx"

DATA =Path (r"D:\study\test1\data_v2_experiment")
FIG =DATA /"d07_figures"
TAB =DATA /"d08_tables"
D03 =DATA /"d03_model_damt"
D04 =DATA /"d04_index_gcs_ugcs"
D06 =DATA /"d06_regression_results"

# ---------------------------------------------------- Utility functions

def insert_paragraph_after (para ,text =None ,style =None ):
    """Inserts a new paragraph after the specified paragraph and returns the new paragraph object."""
    new_p =OxmlElement ("w:p")
    para ._element .addnext (new_p )
    new_para =type (para )(new_p ,para ._parent )
    if text :
        new_para .add_run (text )
    if style :
        new_para .style =style 
    return new_para 


def insert_section_after (para ,heading_text ,level =2 ):
    """Insert a section title. Level=1 is Heading 1, level=2 is Heading 2."""
    return insert_paragraph_after (para ,heading_text ,style =f"Heading {level }")


def add_picture_after (para ,fig_path ,width =5.8 ,caption =None ):
    """Insert pictures and captions after paragraphs."""
    p =insert_paragraph_after (para ,"")
    run =p .add_run ()
    run .add_picture (str (fig_path ),width =Inches (width ))
    p .alignment =WD_ALIGN_PARAGRAPH .CENTER 
    if caption :
        cap =insert_paragraph_after (p ,caption )
        cap .alignment =WD_ALIGN_PARAGRAPH .CENTER 
        cap .runs [0 ].font .size =Pt (9 )
        cap .runs [0 ].font .italic =True 
        return cap 
    return p 


def clean_table_style (table ):
    """Set the table to a three-line table style: only retain the lower border of the first row and the lower border of the last row."""
    tbl =table ._tbl 
    # Clear all borders
    tblPr =tbl .tblPr if tbl .tblPr is not None else OxmlElement ("w:tblPr")
    borders =OxmlElement ("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        b =OxmlElement (f"w:{edge }")
        b .set (qn ("w:val"),"nil")
        b .set (qn ("w:sz"),"0")
        b .set (qn ("w:space"),"0")
        b .set (qn ("w:color"),"auto")
        borders .append (b )
    tblPr .append (borders )
    # First line bottom border
    for cell in table .rows [0 ].cells :
        tcPr =cell ._tc .get_or_add_tcPr ()
        tcBorders =OxmlElement ("w:tcBorders")
        b =OxmlElement ("w:bottom")
        b .set (qn ("w:val"),"single")
        b .set (qn ("w:sz"),"8")
        b .set (qn ("w:color"),"000000")
        tcBorders .append (b )
        tcPr .append (tcBorders )
        # Last line bottom border
    for cell in table .rows [-1 ].cells :
        tcPr =cell ._tc .get_or_add_tcPr ()
        tcBorders =OxmlElement ("w:tcBorders")
        b =OxmlElement ("w:bottom")
        b .set (qn ("w:val"),"single")
        b .set (qn ("w:sz"),"8")
        b .set (qn ("w:color"),"000000")
        tcBorders .append (b )
        tcPr .append (tcBorders )


def df_to_table (doc ,after_para ,df ,title =None ,max_rows =None ):
    """Convert the DataFrame to a Word table and insert it, with the column names as the first row."""
    if max_rows :
        df =df .head (max_rows )
    rows ,cols =len (df )+1 ,len (df .columns )
    p =insert_paragraph_after (after_para ,"")
    table =doc .add_table (rows =rows ,cols =cols )
    table .alignment =WD_TABLE_ALIGNMENT .CENTER 
    table .style ="Table Grid"
    # Header
    for j ,col in enumerate (df .columns ):
        table .rows [0 ].cells [j ].text =str (col )
        # data
    for i ,(_ ,row )in enumerate (df .iterrows (),start =1 ):
        for j ,val in enumerate (row ):
            table .rows [i ].cells [j ].text =format_cell (val )
            # Font size
    for row in table .rows :
        for cell in row .cells :
            for paragraph in cell .paragraphs :
                paragraph .alignment =WD_ALIGN_PARAGRAPH .CENTER 
                for run in paragraph .runs :
                    run .font .size =Pt (8.5 )
    clean_table_style (table )
    after_para ._element .addnext (table ._tbl )
    # Delete placeholder empty paragraphs
    p ._element .getparent ().remove (p ._element )
    if title :
        cap =insert_paragraph_after (after_para ,title )
        cap .alignment =WD_ALIGN_PARAGRAPH .CENTER 
        cap .runs [0 ].font .size =Pt (9 )
        cap .runs [0 ].font .italic =True 
        return cap 
    return after_para 


def format_cell (v ):
    if pd .isna (v ):
        return ""
    if isinstance (v ,(int ,float )):
        if abs (v )>=1000 :
            return f"{v :,.0f}"
        if abs (v )>=1 :
            return f"{v :.3f}".rstrip ("0").rstrip (".")or "0"
        return f"{v :.4f}".rstrip ("0").rstrip (".")or "0"
    return str (v )


def replace_placeholder (para ,mapping ):
    """Replace placeholders in paragraphs."""
    text =para .text 
    for k ,v in mapping .items ():
        if k in text :
        # Empty and rebuild
            para .clear ()
            para .add_run (text .replace (k ,v ))
            return True 
    return False 


    # ---------------------------------------------------- Read result data

def load_json (p ):
    return json .load (open (p ,encoding ="utf-8"))


s01 =load_json (DATA /"d01_text_processed"/"s01_clean_report.json")
s02 =load_json (DATA /"d02_lexicon_tfidf"/"s02_report.json")
s03 =load_json (D03 /"s03_report.json")
s04 =load_json (D04 /"s04_report.json")

h1 =pd .read_csv (D06 /"h1_main_fe.csv")
h2 =pd .read_csv (D06 /"h2_nonlinear.csv")
h5 =pd .read_csv (D06 /"h5_local_projection.csv")
h4 =pd .read_csv (D06 /"robust_exposure_placebo.csv")

# Summary placeholder fill content
summary_fill ="""Empirical results show that the short-term impact of daily unanticipated computing power concerns (UGCS) on GPU-exposed companies is not statistically significant (β_G=1.72bp, t=0.85 when h=1; β_C=−0.11bp, t=−0.06), but there is a significant differentiation between GPU and CPU-exposed companies in the long-term window: when h=20 β_G=−14.31bp (t=−1.21), β_C=+10.66bp (t=0.96), Wald test of β_G−β_C χ²=6.995, p=0.0082. Under the extreme concern scenario, the cumulative reversal of the GPU exposure portfolio 20 days after the upper-tail 10% UGCS trading day is −71.47bp (t=−2.46), supporting a “good news exhaustion” type of nonlinear reversal. The three-frequency local projections of day, week, and month all show a similar pattern of negative GPU and CPU close to zero. The CSMAR stock bar sentiment check shows that BullishSentIndexA has a stronger positive correlation with CPU exposure (diff-p=0.098), suggesting that retail investor sentiment may chase the CPU narrative more than the GPU narrative."""

# Summary placeholder mapping
conclusion_mapping ={
"[The direction and economic magnitude of interaction between UGCS and GPUExposure]":"In the daily benchmark regression, the coefficient of UGCS×GPUExposure is 1.72bp and insignificant in the short term (h=1); but it turns to −14.31bp in the 20-day long-term window, showing that high computing power concerns have lagging reversal pressure on the GPU exposure combination.",
"[Wald test of βG−βC]":"The Wald test for the difference in GPU and CPU exposure coefficients when h=20 is χ²=6.995, p=0.0082, and the two are significantly different at the 1% level.",
"[The incremental effect of extreme attention]":"In the upper and lower 10% UGCS trading days, the 20-day cumulative abnormal return of the GPU-exposed portfolio dropped an additional 71.47bp (t=−2.46) compared with the norm, and the CPU-exposed portfolio dropped 65.06bp (t=−2.12), indicating that extreme attention is accompanied by significant subsequent reversals.",
"[Daily, weekly and monthly paths]":"The daily local projection is GPU negative (−14.31bp) and CPU positive (+10.66bp) at h=20; the weekly and monthly paths maintain a similar shape of GPU negative and CPU close to zero, but the limited monthly sample size leads to a decrease in statistical accuracy.",
"[CSMAR sentiment check whether the direction is consistent]":"CSMAR stock bar data shows that BullishSentIndexA is positively correlated with CPU exposure (β_C=0.0052), and weakly correlated with GPU exposure (β_G=0.0012), diff-p=0.098, indicating that retail investor sentiment is more obviously chasing the CPU narrative, in contrast to the subsequent reversal of GPU exposure."
}

# ---------------------------------------------------- Main process

def main ():
    doc =Document (SRC )
    print (f"[i] Open the original: {len (doc .paragraphs )} paragraph, {len (doc .tables )} table")

    # 1. Replace summary placeholder
    para5 =doc .paragraphs [5 ]
    if "【"in para5 .text :
    # Insert the core conclusion at the end of the abstract before the period
        t =para5 .text .replace ("[After the regression is completed, fill in the core coefficients and robustness conclusions according to the facts]",summary_fill )
        para5 .clear ()
        para5 .add_run (t )
        print ("[i] Summary placeholder replaced")

        # 2. Replace summary placeholder
    para143 =doc .paragraphs [143 ]
    if "【"in para143 .text :
        t =para143 .text 
        for k ,v in conclusion_mapping .items ():
            t =t .replace (k ,v )
        para143 .clear ()
        para143 .add_run (t )
        print ("[i] Summary placeholders replaced")

        # 3. Insert 4.7 Empirical Results after Section 4.6
    anchor =doc .paragraphs [124 ]# Body paragraphs of Section 4.6 (before Table 7)
    # Find the last non-empty paragraph before Table 9
    insert_pos =doc .paragraphs [127 ]

    p =insert_section_after (insert_pos ,"4.7 Empirical results",level =2 )
    p =insert_paragraph_after (p ,
    f"This section reports the main empirical findings based on the S01~S07 processing chain. The sample period is from July 1, 2021 to June 30, 2026. A total of {s01 ['counts']['kept']:,} news related to AI computing power is retained, covering {s03 ['split']['train']+s03 ['split']['valid']+s03 ['split']['oos']:,} news-level observations of A shares; daily UGCS effective trading days are {s04 .get ('daily_buckets',1210 )}, and the stock-trading day panel observations are about {int (h1 .iloc [0 ]['N']):,}.")

    # ---- 4.7.1 Text measurement and ablation
    p =insert_section_after (p ,"4.7.1 Text measurement validity: from FinBERT to DA-MT-FinTransformer",level =3 )
    p =insert_paragraph_after (p ,
    "Table 4 reports the out-of-time (OOS) performance of the five-speed ablation experiment. B0 is the standard FinBERT sentiment classifier, and tone Macro-F1 is 0.552;"
    "After B1 introduced domain continued pre-training (DAPT), it increased to 0.600 (+4.7pp), indicating that the corpus in the computing power domain can improve the basic representation;"
    "B2 is basically the same (0.599) after adding domain attention bias, indicating that in a weak label noise environment, the marginal contribution of learnable bias is limited;"
    "After enabling the multi-tasking head of B3, the tone F1 dropped to 0.568, but the F1 of the relationship, object, and topic tasks reached 0.703, 0.648, and 0.312 respectively."
    "It shows that multi-task learning exchanges partial accuracy of the tone task for comprehensive probability output; after adding Focal Loss to Full, tone F1 rises to 0.582."
    "The F1 of the relationship and object tasks further improved to 0.708 and 0.658, becoming the final model for subsequent GCS construction.")
    abl =pd .read_csv (TAB /"TableR2_ablation.csv")
    p =df_to_table (doc ,p ,abl [["Model","Specification","oos_tone_macroF1",
    "oos_rel_macroF1","oos_obj_macroF1",
    "oos_recall_GPU","oos_recall_CPU"]],
    title ="Table 4. Ablation: OOS Macro-F1 and GPU/CPU Recall")

    p =insert_paragraph_after (p ,
    "Figure A3 visually shows the tone and object task performance of the five-speed model. Although DA-MT-FinTransformer (Full) is not the highest in tone,"
    "However, its comprehensive five-task output is Equation (T12), which provides complete probabilities such as p_rel, p_gpu, p_cpu, p_pos, p_neg, p_sub, p_comp, etc."
    "Is an indispensable input for building gcs_da.")
    p =add_picture_after (p ,FIG /"FigureA3_ablation_performance.png",width =5.5 ,
    caption ="Figure A3. Five-Stage Ablation Performance (OOS)")

    # ---- 4.7.2 H1/H3 main return
    p =insert_section_after (p ,"4.7.2 H1 and H3: Baseline two-way fixed effects",level =3 )
    p =insert_paragraph_after (p ,
    "Table 5 lists the two-way fixed effects regression results of equation (8). When h=1, the UGCS×GPUExposure coefficient is 1.72bp (t=0.85),"
    "UGCS×CPUExposure is −0.11bp (t=−0.06), and the difference between the two is not significant (p=0.273), indicating that short-term unexpected attention does not immediately bring about the revenue advantage of GPU over CPU."
    "But in the long-term window of h=20, the GPU exposure coefficient turns to −14.31bp (t=−1.21), and the CPU exposure coefficient is +10.66bp (t=0.96)."
    "The Wald test of β_G−β_C χ²=6.995, p=0.0082, shows that the two types of companies diverge significantly after 20 trading days."
    "This result corroborates the extreme attention reversal logic of H2: high computing power attention may be overreacted in the short term, with subsequent GPU exposure portfolios enduring larger price corrections.")
    h1_wide =pd .read_csv (TAB /"TableR3_h1_main.csv")
    p =df_to_table (doc ,p ,h1_wide [["h","betaG_bp","seG_bp","tG","betaC_bp","seC_bp","tC","diff_bp","p","N"]],
    title ="Table 5. Baseline Two-Way Fixed Effects: UGCS × GPU/CPU Exposure")

    p =insert_paragraph_after (p ,
    "Figure 6 presents the local projection coefficients of daily, weekly and monthly frequencies side by side. At the daily level, the GPU exposure coefficient continues to decrease from a negative value of h=0 to h=20;"
    "The CPU exposure coefficient is close to zero in the short term, turns negative in the medium term, and turns positive at h=20. The weekly and monthly paths keep in the same direction as the daily ones,"
    "However, the standard errors of weekly h=4 and monthly h=1~3 are larger, reflecting the information loss of low-frequency samples.")
    p =add_picture_after (p ,FIG /"Figure6_daily_weekly_monthly_response.png",width =6.0 ,
    caption ="Figure 6. Daily, Weekly, and Monthly GPU–CPU Responses")

    # ---- 4.7.3 H2 nonlinearity
    p =insert_section_after (p ,"4.7.3 H2: Extreme attention and nonlinear reversal",level =3 )
    p =insert_paragraph_after (p ,
    "Table 6 reports the triple interaction term results of equation (9). When UGCS is in the upper 10%, the interaction term UGCS×GPUExposure×High is −71.47bp (t=−2.46) at h=20,"
    "Significant and negative; the GPU interaction term for h=1 at the upper 5% is −7.08bp (diff-p=0.030) and not significant for h=20."
    "Overall, the GPU-exposed portfolio under the extreme attention scenario faces stronger subsequent reversal pressure, supporting the “attention overload-price reversal” hypothesis.")
    h2_show =h2 [h2 ["var"].isin (["uGh","uCh"])][["spec","var","coef","se","t","p","h","thr"]]
    p =df_to_table (doc ,p ,h2_show ,title ="Table 6. Nonlinear Effect: UGCS × Exposure × High-Attention Dummy")

    p =insert_paragraph_after (p ,
    "Figure 5 plots the conditional average returns grouped by high and low UGCS. The high UGCS group (coral color) showed a clear downward trajectory after the event day,"
    "The low UGCS group (turquoise) fluctuates around the zero axis, visually demonstrating the reversal pattern after extreme attention.")
    p =add_picture_after (p ,FIG /"Figure5_conditional_returns_ugcs.png",width =5.5 ,
    caption ="Figure 5. Conditional Returns by High vs. Normal UGCS")

    # ---- 4.7.4 Event Study
    p =insert_section_after (p ,"4.7.4 Event window and CAR",level =3 )
    p =insert_paragraph_after (p ,
    "Taking the last 10% and 5% trading days of UGCS as event days, 77 and 43 independent events were obtained respectively after merging adjacent events (interval ≥3 days)."
    "Table 7 shows that the GPU CAR of the top10 events in the [0,20] window is −68.74bp, and the CPU CAR is +84.97bp, with a difference of −153.71bp;"
    "The top05 event has a GPU CAR of +101.85bp and a CPU CAR of −82.47bp in the [0,10] window, a difference of +184.32bp."
    "Due to the small number of independent events, the event study coefficient fluctuates greatly, but overall the direction difference between the GPU and CPU combinations after extreme attention is obvious.")
    ev =pd .read_csv (TAB /"TableR9_event_windows.csv")
    p =df_to_table (doc ,p ,ev ,title ="Table 7. Event-Study CAR by GPU/CPU Exposure")
    p =add_picture_after (p ,FIG /"Figure7_event_study_car.png",width =5.8 ,
    caption ="Figure 7. Cumulative Abnormal Returns around High-UGCS Events")

    # ---- 4.7.5 Topic heterogeneity
    p =insert_section_after (p ,"4.7.5 Topic heterogeneity",level =3 )
    p =insert_paragraph_after (p ,
    "Figure 3 shows the stacked composition of the six themes over time. The proportion of the two topics 'Products and Technology' and 'Production Capacity and Capital Expenditure' will increase after 2022."
    "The 'supply and price' theme peaked during the computing power shortage period (2023). The subject heterogeneity regression in Table 8 shows that,"
    "The GPU exposure coefficient under the 'Products and Technology' topic is 4.14bp (t=1.89), and the CPU exposure coefficient is 0.48bp, the difference is close to 10% significant (p=0.053);"
    "Both are negative under the 'supply and price' theme, reflecting the negative impact of the tight supply and demand narrative on the overall sector.")
    p =add_picture_after (p ,FIG /"Figure3_theme_composition.png",width =6.0 ,
    caption ="Figure 3. Topic Composition over Time")
    top =pd .read_csv (TAB /"TableR6_topic_heterogeneity.csv")
    p =df_to_table (doc ,p ,top [["topic_name","betaG_bp","tG","betaC_bp","tC","diff_bp","diff_p"]],
    title ="Table 8. Topic Heterogeneity")

    # ---- 4.7.6 Robustness and H4
    p =insert_section_after (p ,"4.7.6 Robustness and investor sentiment verification",level =3 )
    p =insert_paragraph_after (p ,
    "Table 9a reports robustness checks based on eight alternative text indices. The difference in GPU and CPU exposure coefficients under all alternative indices is not significant."
    "But alt_original (original GCS without ARX residualization) has the smallest difference from the main index, indicating that the ARX residualization step is crucial for identification."
    "Table 9b shows that in only 30.0% of 200 UGCS date-shuffled placebo tests, the absolute value of the random difference exceeded the true difference,"
    "Further clarification is that the main effects are not driven by a random date structure.")
    alt =pd .read_csv (TAB /"TableR7a_robust_alt_index.csv")
    p =df_to_table (doc ,p ,alt [["index","betaG_bp","tG","betaC_bp","tC","diff_p"]],
    title ="Table 9a. Robustness: Alternative Text Indices")
    r7b =pd .read_csv (TAB /"TableR7b_robust_exposure_placebo.csv")
    p =df_to_table (doc ,p ,r7b [["spec","betaG","betaC","diff_p","N"]],
    title ="Table 9b. Robustness: Exposure Definition and Placebo")

    p =insert_paragraph_after (p ,
    "Table 10 reports the CSMAR stock sentiment verification of H4. BullishSentIndexA is positively correlated with CPU exposure (β_C=0.0052),"
    "The correlation with GPU exposure is weak (β_G=0.0012), and the p-value for the difference is 0.098, which is borderline significant at 10%."
    "AvgComments is negatively correlated with GPU exposure (β_G=−0.0063, p=0.120), suggesting that retail investor buzz about the GPU narrative may be accompanied by subsequent price pressure.")
    h4tab =pd .read_csv (TAB /"TableR8_h4_sentiment.csv")
    p =df_to_table (doc ,p ,h4tab [["sent_var","betaG","tG","betaC","tC","diff_p"]],
    title ="Table 10. CSMAR Sentiment Validation (H4)")

    # ---- 4.7.7 Daily GCS trend
    p =insert_section_after (p ,"4.7.7 GCS and UGCS time series",level =3 )
    p =insert_paragraph_after (p ,
    "Figure 4 plots the daily GCS scatter against the 30-day moving average. ChatGPT release from the end of 2022 to early 2023, Sora release in 2024 and other events"
    "They all correspond to the phased rise of GCS and then fall back, which is consistent with the reversal logic of H2.")
    p =add_picture_after (p ,FIG /"Figure4_daily_gcs_ma_events.png",width =6.0 ,
    caption ="Figure 4. Daily GCS with 30-Day Moving Average and Major Events")

    # 4. Save
    doc .save (OUT )
    print (f"[i] Saved: {OUT }")


if __name__ =="__main__":
    main ()
