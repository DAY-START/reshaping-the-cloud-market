# -*- coding: utf-8 -*-

# ============================================================================
# RESTRICTED USE -- All rights reserved.
# The data, experimental results, and source code in this repository are
# provided for archival / reproducibility reference only. No part of the
# data, experiments, or code may be used, copied, modified, redistributed,
# or incorporated into other works without the author's explicit written
# permission. See LICENSE and NOTICE.md.
# ============================================================================

"""v6_s10_process_doc.py
=====================
Generate a stand-alone Word document describing the experimental procedure:
- Data sources and directory specifications
- What each step of the script does, key parameters, and time consumption
- Description of key intermediate products
- Summary of final results and main conclusions"""
import json 
import io 
import sys 
from pathlib import Path 
from datetime import datetime 
from docx import Document 
from docx .shared import Inches ,Pt ,RGBColor 
from docx .enum .text import WD_ALIGN_PARAGRAPH 
from docx .enum .table import WD_TABLE_ALIGNMENT 
from docx .oxml import OxmlElement 
from docx .oxml .ns import qn 
import pandas as pd 

sys .stdout =io .TextIOWrapper (sys .stdout .buffer ,encoding ="utf-8",errors ="replace")

OUT =r"C:\Users\ADMIN\Desktop\实验过程说明_AI算力替代与股价分化_V6.docx"
DATA =Path (r"D:\study\test1\data_v2_experiment")
ROOT =Path (r"D:\study\test1")

s01 =json .load (open (DATA /"d01_text_processed"/"s01_clean_report.json",encoding ="utf-8"))
s02 =json .load (open (DATA /"d02_lexicon_tfidf"/"s02_report.json",encoding ="utf-8"))
s03 =json .load (open (DATA /"d03_model_damt"/"s03_report.json",encoding ="utf-8"))
s04 =json .load (open (DATA /"d04_index_gcs_ugcs"/"s04_report.json",encoding ="utf-8"))


def add_heading (doc ,text ,level =1 ):
    return doc .add_heading (text ,level =level )


def add_para (doc ,text ,bold =False ):
    p =doc .add_paragraph ()
    r =p .add_run (text )
    r .bold =bold 
    return p 


def clean_table_style (table ):
    tbl =table ._tbl 
    for row in table .rows :
        for cell in row .cells :
            tcPr =cell ._tc .get_or_add_tcPr ()
            tcBorders =OxmlElement ("w:tcBorders")
            for edge in ("top","left","bottom","right"):
                b =OxmlElement (f"w:{edge }")
                b .set (qn ("w:val"),"nil")
                tcBorders .append (b )
            tcPr .append (tcBorders )
            # First line underline
    for cell in table .rows [0 ].cells :
        tcPr =cell ._tc .get_or_add_tcPr ()
        tcBorders =OxmlElement ("w:tcBorders")
        b =OxmlElement ("w:bottom")
        b .set (qn ("w:val"),"single")
        b .set (qn ("w:sz"),"8")
        b .set (qn ("w:color"),"000000")
        tcBorders .append (b )
        tcPr .append (tcBorders )
        # Last line underline
    for cell in table .rows [-1 ].cells :
        tcPr =cell ._tc .get_or_add_tcPr ()
        tcBorders =OxmlElement ("w:tcBorders")
        b =OxmlElement ("w:bottom")
        b .set (qn ("w:val"),"single")
        b .set (qn ("w:sz"),"8")
        b .set (qn ("w:color"),"000000")
        tcBorders .append (b )
        tcPr .append (tcBorders )


def df_to_docx (doc ,df ):
    table =doc .add_table (rows =len (df )+1 ,cols =len (df .columns ))
    table .alignment =WD_TABLE_ALIGNMENT .CENTER 
    for j ,c in enumerate (df .columns ):
        table .rows [0 ].cells [j ].text =str (c )
    for i ,(_ ,row )in enumerate (df .iterrows (),start =1 ):
        for j ,v in enumerate (row ):
            table .rows [i ].cells [j ].text =str (v )
    for row in table .rows :
        for cell in row .cells :
            for p in cell .paragraphs :
                p .alignment =WD_ALIGN_PARAGRAPH .CENTER 
                for r in p .runs :
                    r .font .size =Pt (9 )
    clean_table_style (table )


def main ():
    doc =Document ()
    doc .add_heading ("Experimental process documentation",level =0 ).alignment =WD_ALIGN_PARAGRAPH .CENTER 
    add_para (doc ,f"Thesis title: Reshaping the cloud market: How AI computing power substitution drives stock price differentiation (V6 multi-frequency and DA-Transformer enhanced version)")
    add_para (doc ,f"Generation time: {datetime .now ().strftime ('%Y-%m-%d %H:%M')}")
    add_para (doc ,f"Execution environment: Python 3.11 / Windows 10 / CPU training")

    # 1. Directory and naming conventions
    add_heading (doc ,"1 Directory and script naming convention",level =1 )
    add_para (doc ,"原始数据保留在原位，实验中间产物统一存放在 D:\\study\\test1\\data_v2_experiment\\，脚本统一存放在 D:\\study\\test1\\scripts_v6_pipeline\\。")
    tbl_data =[
    ["Table of contents","use"],
    ["data_v2_experiment/d01_text_processed","Cleaned news metadata and text"],
    ["data_v2_experiment/d02_lexicon_tfidf","Domain dictionary, TF-IDF, weak tags and lexicon score"],
    ["data_v2_experiment/d03_model_damt","DA-MT-FinTransformer encoding, weight, ablation index and full-sample inference score"],
    ["data_v2_experiment/d04_index_gcs_ugcs","Daily/weekly/monthly GCS, UGCS, subject and alternative caliber indexes"],
    ["data_v2_experiment/d05_panel_dataset","Stock panel, ex-ante exposure, firm master"],
    ["data_v2_experiment/d06_regression_results","H1~H5, robustness, placebo, sentiment verification regression table"],
    ["data_v2_experiment/d07_figures","300 dpi top issue color graphics"],
    ["data_v2_experiment/d08_tables","Summary table CSV"],
    ["data_v2_experiment/d09_logs","Operation log of each step"],
    ]
    table =doc .add_table (rows =len (tbl_data ),cols =2 )
    for i ,row in enumerate (tbl_data ):
        for j ,v in enumerate (row ):
            table .rows [i ].cells [j ].text =v 
    clean_table_style (table )

    add_heading (doc ,"2 Script list and running sequence",level =1 )
    steps =[
    ["v6_s01_news_preprocess.py","News cleaning, deduplication, date normalization, text placement","206s"],
    ["v6_s02_lexicon_tfidf.py","jieba word segmentation, domain dictionary hits, TF-IDF, weak tags and gcs_lex","881s"],
    ["v6_s03_damt_transformer.py","DAPT continues pre-training, five-level ablation, full-sample inference gcs_da","Approximately 90 minutes in stages"],
    ["v6_s04_gcs_ugcs.py","In-source rolling normalization, strictly rolling ARX residualization, tri-frequency index","140s"],
    ["v6_s05_panel_exposure.py","Stock panels, rolling beta, ex-ante exposure","37s"],
    ["v6_s06_regressions.py","H1~H5 all regression, robustness, placebo, sentiment verification","176s"],
    ["v6_s07_figures_tables.py","Top issue color graphics and summary table","8s"],
    ["v6_s08_docx_scan.py","Scan paper docx structure and placeholders","<1s"],
    ["v6_s09_fill_paper.py","Populate results chart into paper docx","<1s"],
    ["v6_s10_process_doc.py","Generate this documentation","<1s"],
    ]
    table =doc .add_table (rows =len (steps )+1 ,cols =3 )
    hdr =["script name","Function","time consuming"]
    for j ,v in enumerate (hdr ):
        table .rows [0 ].cells [j ].text =v 
    for i ,row in enumerate (steps ,start =1 ):
        for j ,v in enumerate (row ):
            table .rows [i ].cells [j ].text =v 
    clean_table_style (table )

    add_heading (doc ,"3 Data sources and original data retention",level =1 )
    add_para (doc ,"原始数据未删除，仍保留在 D:\\study\\test1\\ 根目录下：")
    add_para (doc ,f"• News original data: news_data_v1_initial\\, a total of {s01 ['counts']['total']:,} original news;")
    add_para (doc ,"• 股票原始数据：stock_data_v1_initial\\all_66stock_final.csv，66 只股票日线；")
    add_para (doc ,"• 情绪原始数据：investor_sentiment_data_v1_initial\\total_sentiment_clean.csv，约 780 万行股吧数据；")
    add_para (doc ,"• iFind and References: iffind_v1_initial, references_v1_initial.")

    add_heading (doc ,"4 Key Steps Instructions",level =1 )

    add_heading (doc ,"4.1 News Cleaning (S01)",level =2 )
    add_para (doc ,f"Starting from {s01 ['counts']['total']:,} original news items, after going through AI computing power filtering, deduplication, text length and language verification, date analysis and other steps, {s01 ['counts']['kept']:,} items are finally retained. Key reasons for elimination: non-AI computing power {s01 ['counts']['drop_not_ai_compute']:,}, duplicate title date {s01 ['counts']['drop_title_date_dup']:,}, text too short {s01 ['counts']['drop_short']:,}, market chat {s01 ['counts']['drop_market_talk']:,}.")

    add_heading (doc ,"4.2 Dictionary method and weak tags (S02)",level =2 )
    rlt_labeled =s02 ['label_distribution']['rlt0']+s02 ['label_distribution']['rlt1']+s02 ['label_distribution']['rlt2']+s02 ['label_distribution']['rlt3']
    top_labeled =sum (v for k ,v in s02 ['label_distribution'].items ()if k .startswith ('top')and k !='top_unlabeled')
    add_para (doc ,f"Build domain dictionaries such as GPU/CPU/relationship/event/intonation/intensity/topic, calculate TF-IDF weights and generate five-task weak labels. There are {rlt_labeled :,} valid labels for the relationship class and {top_labeled :,} valid labels for the topic class; unlabeled samples are masked by the -1 label during training.")

    add_heading (doc ,"4.3 DA-MT-FinTransformer（S03）",level =2 )
    add_para (doc ,f"Model structure: L_SEQ={s03 ['arch']['L_SEQ']}, D_MODEL={s03 ['arch']['D_MODEL']}, N_LAYER={s03 ['arch']['N_LAYER']}, N_HEAD={s03 ['arch']['N_HEAD']}, D_FF={s03 ['arch']['D_FF']}.")
    add_para (doc ,f"Training settings: DAPT {s03 ['dapt']['steps']} steps (loss {s03 ['dapt']['loss_first']} → {s03 ['dapt']['loss_last']}, final PPL {s03 ['dapt']['final_ppl']}), fine-tuning epoch={s03 ['train_cfg']['ft_epoch']}, learning rate {s03 ['train_cfg']['lr']}.")
    add_para (doc ,"Key optimizations: staged breakpoint continued training, sparse MLM head, Pre-LN residual, category weighted CE, warmup+cosine scheduling to avoid early model degradation into majority class prediction.")
    add_para (doc ,"Five levels of ablation (OOS tone Macro-F1): B0=0.552, B1=0.600, B2=0.599, B3=0.568, Full=0.582.")
    add_para (doc ,f"Full sample inference: {s03 ['n_docs']:,} news, gcs_da mean {s03 ['gcs_da_stat']['mean']}, standard deviation {s03 ['gcs_da_stat']['sd']}.")

    add_heading (doc ,"4.4 GCS/UGCS tri-frequency structure (S04)",level =2 )
    add_para (doc ,f"Valid trading days are {s04 ['daily']['n_buckets']} per day, {s04 ['weekly']['n_buckets']} per week, and {s04 ['monthly']['n_buckets']} per month. Day/week/month ARX R² are {s04 ['daily']['arx_mean_R2']:.3f}/{s04 ['weekly']['arx_mean_R2']:.3f}/{s04 ['monthly']['arx_mean_R2']:.3f} respectively.")

    add_heading (doc ,"4.5 Panels and Exposure (S05)",level =2 )
    add_para (doc ,"By scanning the news text according to the company names and aliases of 66 stocks, 41,863 company-news co-occurrences were obtained;"
    "Textual alternative caliber ex ante exposures were constructed with a 250-day rolling window and z-normalized to the cross-section.")

    add_heading (doc ,"4.6 Regression and hypothesis testing (S06)",level =2 )
    add_para (doc ,"All regressions use two-way fixed effects (firm + time) + two-way clustered robust standard errors.")
    h1 =pd .read_csv (DATA /"d08_tables"/"TableR3_h1_main.csv")
    add_para (doc ,f"H1/H3: When h=20, β_G={h1 [h1 .h ==20 ].iloc [0 ]['betaG_bp']:.2f}bp, β_C={h1 [h1 .h ==20 ].iloc [0 ]['betaC_bp']:.2f}bp, Wald p={h1 [h1 .h ==20 ].iloc [0 ]['p']:.4f}.")
    h4 =pd .read_csv (DATA /"d06_regression_results"/"robust_exposure_placebo.csv")
    add_para (doc ,f"Placebo: {h4 [h4 .spec =='placebo_shuffle_ugcs'].iloc [0 ]['diff_p']*100 :.1f}% of random differences exceed true differences in 200 UGCS date shuffles.")

    add_heading (doc ,"5 Main conclusions",level =1 )
    add_para (doc ,"1. The differential impact of daily short-term (h=1) UGCS on GPU/CPU exposed stocks is not significant;"
    "However, significant differentiation occurs in the long-term window of h=20 (Wald p=0.0082 for β_G−β_C).")
    add_para (doc ,"2. There is a significant subsequent reversal in the extreme concern scenario: 20 days after the upper tail 10% UGCS event, the GPU exposure portfolio dropped by approximately 71bp additionally (t=−2.46).")
    add_para (doc ,"3. The daily-week-month three-frequency path maintains a consistent pattern of negative GPU and close to zero CPU, but the monthly sample size is limited.")
    add_para (doc ,"4. CSMAR stock bar sentiment shows that retail investor sentiment is stronger in pursuit of the CPU narrative (BullishSentIndexA diff-p=0.098).")
    add_para (doc ,"5. A robustness test of the alternative text index with 200 placebo tests supports that the main effects are not driven by random structure.")

    add_heading (doc ,"6 Notes and Limitations",level =1 )
    add_para (doc ,"• The empirical results show the differentiation of the long-term negative direction of GPU and the long-term positive direction of CPU, which is opposite to the direction of positive direction of GPU and negative direction of CPU of H1 null hypothesis."
    "It needs to be explained in conjunction with the reversal logic of H2 and the ‘profit-exhausting’ mechanism.")
    add_para (doc ,"• The monthly panel sample size is only about 1,900 items, and the statistical power for low-frequency inference is limited.")
    add_para (doc ,"• Weak labels are based on distant lexicon supervision and contain noise, which may lower the absolute F1 of the multi-task model.")

    doc .save (OUT )
    print (f"[i] Saved: {OUT }")


if __name__ =="__main__":
    main ()
